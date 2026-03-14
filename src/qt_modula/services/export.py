"""High-performance export writer services."""

from __future__ import annotations

import csv
import io
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Protocol

import orjson
import xlsxwriter  # type: ignore[import-untyped]
from docx import Document
from openpyxl import load_workbook  # type: ignore[import-untyped]

from qt_modula.services.errors import ServiceError

_TABLE_FORMATS = {"csv", "jsonl", "xlsx"}
_TEXT_FORMATS = {"txt", "docx", "json"}
_MODES = {"overwrite", "append"}
_JSON_KEY_CONFLICT_POLICIES = {"overwrite", "error", "skip"}
_JSON_DUPLICATE_KEY_POLICIES = {"error", "last_wins"}
_JSON_DUMP_OPTIONS = orjson.OPT_INDENT_2 | orjson.OPT_SORT_KEYS
_EXPORT_FOLDER_LIMIT_BYTES = 100 * 1024 * 1024


@dataclass(frozen=True, slots=True)
class ExportRequest:
    """Typed table export request model."""

    path: Path
    rows: list[dict[str, Any]]
    mode: str = "overwrite"


@dataclass(frozen=True, slots=True)
class ExportResult:
    """Typed table export response model."""

    path: Path
    row_count: int
    total_row_count: int


@dataclass(frozen=True, slots=True)
class TextExportRequest:
    """Typed text export request model."""

    path: Path
    text: str
    mode: str = "overwrite"
    section_title: str = ""
    json_dictionary_bound: bool = False
    json_key_conflict: str = "overwrite"
    json_duplicate_keys: str = "error"


@dataclass(frozen=True, slots=True)
class TextExportResult:
    """Typed text export response model."""

    path: Path
    char_count: int
    line_count: int
    wrote: bool = True


class ExportWriter(Protocol):
    """Export writer protocol for tabular datasets."""

    def write(self, request: ExportRequest) -> ExportResult:
        """Write rows to path and return metadata."""


class TextExportWriter(Protocol):
    """Export writer protocol for text payloads."""

    def write(self, request: TextExportRequest) -> TextExportResult:
        """Write text to path and return metadata."""


def _normalized_mode(mode: str) -> str:
    token = mode.strip().lower()
    if token in _MODES:
        return token
    raise ServiceError(
        kind="validation",
        message=f"Unsupported export mode '{mode}'.",
        provider="export",
    )


def _normalized_json_key_conflict_policy(policy: str) -> str:
    token = policy.strip().lower()
    if token in _JSON_KEY_CONFLICT_POLICIES:
        return token
    raise ServiceError(
        kind="validation",
        message=f"Unsupported JSON key conflict policy '{policy}'.",
        provider="export",
    )


def _normalized_json_duplicate_key_policy(policy: str) -> str:
    token = policy.strip().lower()
    if token in _JSON_DUPLICATE_KEY_POLICIES:
        return token
    raise ServiceError(
        kind="validation",
        message=f"Unsupported JSON duplicate-key policy '{policy}'.",
        provider="export",
    )


def _coerce_rows(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    coerced: list[dict[str, Any]] = []
    for row in rows:
        coerced.append({str(key): value for key, value in row.items()})
    return coerced


def _sorted_headers(rows: list[dict[str, Any]]) -> list[str]:
    return sorted({key for row in rows for key in row})


def _csv_payload(rows: list[dict[str, Any]]) -> bytes:
    headers = _sorted_headers(rows)
    buffer = io.StringIO(newline="")
    writer = csv.DictWriter(buffer, fieldnames=headers)
    writer.writeheader()
    writer.writerows(rows)
    return buffer.getvalue().encode("utf-8")


def _read_csv(path: Path) -> list[dict[str, Any]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8", newline="") as handle:
        reader = csv.DictReader(handle)
        rows: list[dict[str, Any]] = []
        for item in reader:
            row: dict[str, Any] = {}
            for key, value in item.items():
                if key is None:
                    continue
                row[str(key)] = value
            rows.append(row)
        return rows


def _jsonl_payload(rows: list[dict[str, Any]]) -> bytes:
    payload = bytearray()
    for row in rows:
        payload.extend(orjson.dumps(row, option=orjson.OPT_SORT_KEYS))
        payload.extend(b"\n")
    return bytes(payload)


def _temporary_path(path: Path) -> Path:
    return path.with_name(f".{path.name}.tmp")


def _folder_size(path: Path, *, exclude: set[Path] | None = None) -> int:
    if not path.exists():
        return 0
    excluded = {(item.resolve() if item.exists() else item) for item in (exclude or set())}
    total = 0
    for entry in path.rglob("*"):
        if not entry.is_file():
            continue
        resolved = entry.resolve()
        if resolved in excluded or entry in excluded:
            continue
        total += entry.stat().st_size
    return total


def _format_size(num_bytes: int) -> str:
    size = float(num_bytes)
    for unit in ("bytes", "KB", "MB", "GB"):
        if size < 1024 or unit == "GB":
            if unit == "bytes":
                return f"{int(size)} {unit}"
            return f"{size:.1f} {unit}"
        size /= 1024
    return f"{size:.1f} GB"


def _enforce_export_folder_limit(
    path: Path,
    *,
    new_file_size: int,
    current_folder_size: int | None = None,
) -> None:
    if current_folder_size is None:
        current_total = _folder_size(path.parent)
    else:
        current_total = current_folder_size
    existing_size = path.stat().st_size if path.exists() else 0
    projected_total = max(0, current_total - existing_size + new_file_size)
    if projected_total <= _EXPORT_FOLDER_LIMIT_BYTES:
        return
    if projected_total < current_total:
        return

    folder_text = str(path.parent)
    limit_text = _format_size(_EXPORT_FOLDER_LIMIT_BYTES)
    if current_total > _EXPORT_FOLDER_LIMIT_BYTES:
        message = (
            f"Export folder limit reached: '{folder_text}' is already using "
            f"{_format_size(current_total)}. Remove older exports before writing more files."
        )
    else:
        message = (
            f"Export folder limit reached: writing '{path.name}' would increase "
            f"'{folder_text}' to {_format_size(projected_total)} (limit: {limit_text}). "
            "Remove older exports and try again."
        )
    raise ServiceError(kind="validation", message=message, provider="export")


def _atomic_write_with_limit(path: Path, payload: bytes) -> None:
    _enforce_export_folder_limit(path, new_file_size=len(payload))
    _atomic_write(path, payload)


def _replace_with_limit(path: Path, tmp_path: Path) -> None:
    current_total = _folder_size(path.parent, exclude={tmp_path})
    _enforce_export_folder_limit(
        path,
        new_file_size=tmp_path.stat().st_size,
        current_folder_size=current_total,
    )
    tmp_path.replace(path)


def _atomic_write(path: Path, payload: bytes) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = _temporary_path(path)
    with tmp.open("wb") as handle:
        handle.write(payload)
        handle.flush()
    tmp.replace(path)


def _count_jsonl_rows(path: Path) -> int:
    if not path.exists():
        return 0
    with path.open("rb") as handle:
        return sum(1 for _ in handle)


def _unique_headers(raw_headers: tuple[Any, ...]) -> list[str]:
    result: list[str] = []
    seen: dict[str, int] = {}
    for idx, value in enumerate(raw_headers):
        token = str(value).strip() if value is not None else ""
        if not token:
            token = f"c{idx}"
        base = token
        suffix = seen.get(base, 0)
        while token in seen:
            suffix += 1
            token = f"{base}_{suffix}"
        seen[base] = suffix
        seen[token] = 0
        result.append(token)
    return result


def _read_xlsx(path: Path) -> list[dict[str, Any]]:
    if not path.exists():
        return []
    workbook = load_workbook(filename=str(path), data_only=True, read_only=True)
    try:
        sheet = workbook.active
        iterator = sheet.iter_rows(values_only=True)
        header_row = next(iterator, None)
        if header_row is None:
            return []
        headers = _unique_headers(tuple(header_row))
        rows: list[dict[str, Any]] = []
        for line in iterator:
            row = {
                headers[index]: line[index] if index < len(line) else None
                for index in range(len(headers))
            }
            if any(value is not None for value in row.values()):
                rows.append(row)
        return rows
    finally:
        workbook.close()


def _write_xlsx(path: Path, rows: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    headers = _sorted_headers(rows)
    workbook = xlsxwriter.Workbook(str(path))
    sheet = workbook.add_worksheet("Sheet1")
    try:
        for col_idx, header in enumerate(headers):
            sheet.write(0, col_idx, header)
        for row_idx, row in enumerate(rows, start=1):
            for col_idx, header in enumerate(headers):
                sheet.write(row_idx, col_idx, row.get(header))
    finally:
        workbook.close()


def _append_text(existing: str, incoming: str, section_title: str) -> str:
    if not existing:
        if section_title:
            return f"{section_title}\n{incoming}"
        return incoming

    base = existing.rstrip("\n")
    if section_title:
        return f"{base}\n\n{section_title}\n{incoming}"
    return f"{base}\n{incoming}"


def _decode_text(payload: bytes) -> str:
    return payload.decode("utf-8")


def _line_count(text: str) -> int:
    return len(text.splitlines()) if text else 0


def _json_stats(payload: bytes) -> tuple[int, int]:
    text = _decode_text(payload)
    return len(text), _line_count(text)


def _parse_json_object(text: str, *, duplicate_policy: str) -> dict[str, Any]:
    duplicate_token = _normalized_json_duplicate_key_policy(duplicate_policy)
    if duplicate_token == "last_wins":
        try:
            loaded = orjson.loads(text)
        except orjson.JSONDecodeError as exc:
            raise ServiceError(
                kind="validation",
                message=f"Invalid JSON payload: {exc}.",
                provider="export",
            ) from exc
        if isinstance(loaded, dict):
            return loaded
        raise ServiceError(
            kind="validation",
            message="JSON Dictionary Bound input must be a JSON object.",
            provider="export",
        )

    duplicates: set[str] = set()

    def _pairs_hook(pairs: list[tuple[str, Any]]) -> dict[str, Any]:
        data: dict[str, Any] = {}
        for key, value in pairs:
            token = str(key)
            if token in data:
                duplicates.add(token)
            data[token] = value
        return data

    try:
        loaded = json.loads(text, object_pairs_hook=_pairs_hook)
    except json.JSONDecodeError as exc:
        raise ServiceError(
            kind="validation",
            message=f"Invalid JSON payload: {exc}.",
            provider="export",
        ) from exc

    if duplicates:
        names = ", ".join(sorted(duplicates))
        raise ServiceError(
            kind="validation",
            message=f"Duplicate JSON keys are not allowed: {names}.",
            provider="export",
        )
    if isinstance(loaded, dict):
        return loaded
    raise ServiceError(
        kind="validation",
        message="JSON Dictionary Bound input must be a JSON object.",
        provider="export",
    )


def _load_json_object(path: Path) -> tuple[dict[str, Any], bytes]:
    payload = path.read_bytes()
    try:
        loaded = orjson.loads(payload)
    except orjson.JSONDecodeError as exc:
        raise ServiceError(
            kind="validation",
            message=f"Existing JSON file is invalid: {exc}.",
            provider="export",
        ) from exc
    if isinstance(loaded, dict):
        return loaded, payload
    raise ServiceError(
        kind="validation",
        message="Existing JSON file must contain a JSON object root.",
        provider="export",
    )


def _dump_json_object(data: dict[str, Any]) -> bytes:
    return orjson.dumps(data, option=_JSON_DUMP_OPTIONS)


def _docx_lines(document: Any) -> list[str]:
    return [str(paragraph.text or "") for paragraph in document.paragraphs]


def _append_docx(document: Any, incoming: str, section_title: str) -> None:
    has_existing_content = any(line for line in _docx_lines(document))
    if section_title:
        if has_existing_content:
            document.add_paragraph("")
            document.add_paragraph("")
        document.add_paragraph(section_title)
    lines = incoming.splitlines() or [""]
    for line in lines:
        document.add_paragraph(line)


class CsvExportWriter:
    """Deterministic CSV export writer with overwrite/append modes."""

    def write(self, request: ExportRequest) -> ExportResult:
        path = request.path
        rows = _coerce_rows(request.rows)
        mode = _normalized_mode(request.mode)

        if mode == "append" and path.exists():
            existing = _read_csv(path)
            combined = existing + rows
        else:
            combined = rows

        _atomic_write_with_limit(path, _csv_payload(combined))
        return ExportResult(
            path=path,
            row_count=len(rows),
            total_row_count=len(combined),
        )


class JsonlExportWriter:
    """Deterministic JSONL export writer with overwrite/append modes."""

    def write(self, request: ExportRequest) -> ExportResult:
        path = request.path
        rows = _coerce_rows(request.rows)
        mode = _normalized_mode(request.mode)
        payload = _jsonl_payload(rows)

        if mode == "append":
            existing_count = _count_jsonl_rows(path)
            existing_payload = path.read_bytes() if path.exists() else b""
            if existing_payload and not existing_payload.endswith(b"\n") and payload:
                final_payload = existing_payload + b"\n" + payload
            else:
                final_payload = existing_payload + payload
            _atomic_write_with_limit(path, final_payload)
            total_count = existing_count + len(rows)
        else:
            _atomic_write_with_limit(path, payload)
            total_count = len(rows)

        return ExportResult(
            path=path,
            row_count=len(rows),
            total_row_count=total_count,
        )


class XlsxExportWriter:
    """XLSX export writer with overwrite/append modes."""

    def write(self, request: ExportRequest) -> ExportResult:
        path = request.path
        rows = _coerce_rows(request.rows)
        mode = _normalized_mode(request.mode)

        if mode == "append" and path.exists():
            existing = _read_xlsx(path)
            combined = existing + rows
        else:
            combined = rows

        path.parent.mkdir(parents=True, exist_ok=True)
        tmp = _temporary_path(path)
        try:
            if tmp.exists():
                tmp.unlink()
            _write_xlsx(tmp, combined)
            _replace_with_limit(path, tmp)
        finally:
            if tmp.exists():
                tmp.unlink()
        return ExportResult(
            path=path,
            row_count=len(rows),
            total_row_count=len(combined),
        )


class TxtTextExportWriter:
    """Deterministic UTF-8 text writer."""

    def write(self, request: TextExportRequest) -> TextExportResult:
        path = request.path
        mode = _normalized_mode(request.mode)
        path.parent.mkdir(parents=True, exist_ok=True)

        incoming = request.text
        section_title = request.section_title.strip()
        if mode == "append":
            existing = path.read_text(encoding="utf-8") if path.exists() else ""
            final_text = _append_text(existing, incoming, section_title)
        else:
            final_text = incoming

        _atomic_write_with_limit(path, final_text.encode("utf-8"))
        return TextExportResult(
            path=path,
            char_count=len(final_text),
            line_count=_line_count(final_text),
        )


class DocxTextExportWriter:
    """Deterministic DOCX text writer using python-docx."""

    def write(self, request: TextExportRequest) -> TextExportResult:
        path = request.path
        mode = _normalized_mode(request.mode)
        path.parent.mkdir(parents=True, exist_ok=True)

        if mode == "append":
            document = Document(str(path)) if path.exists() else Document()
            _append_docx(document, request.text, request.section_title.strip())
        else:
            document = Document()
            lines = request.text.splitlines() or [""]
            for line in lines:
                document.add_paragraph(line)

        tmp = _temporary_path(path)
        try:
            if tmp.exists():
                tmp.unlink()
            document.save(str(tmp))
            _replace_with_limit(path, tmp)
        finally:
            if tmp.exists():
                tmp.unlink()

        snapshot = Document(str(path))
        final_text = "\n".join(_docx_lines(snapshot))
        return TextExportResult(
            path=path,
            char_count=len(final_text),
            line_count=_line_count(final_text),
        )


class JsonTextExportWriter:
    """Deterministic JSON object writer for Text Export workflows."""

    def write(self, request: TextExportRequest) -> TextExportResult:
        path = request.path
        mode = _normalized_mode(request.mode)
        key_conflict = _normalized_json_key_conflict_policy(request.json_key_conflict)
        duplicate_policy = _normalized_json_duplicate_key_policy(request.json_duplicate_keys)
        path.parent.mkdir(parents=True, exist_ok=True)

        if bool(request.json_dictionary_bound):
            incoming = _parse_json_object(request.text, duplicate_policy=duplicate_policy)
        else:
            section_title = request.section_title.strip()
            if not section_title:
                raise ServiceError(
                    kind="validation",
                    message=(
                        "Section Title is required when JSON Dictionary Bound is disabled."
                    ),
                    provider="export",
                )
            incoming = {section_title: request.text}

        if mode == "overwrite":
            payload = _dump_json_object(incoming)
            _atomic_write_with_limit(path, payload)
            char_count, line_count = _json_stats(payload)
            return TextExportResult(
                path=path,
                char_count=char_count,
                line_count=line_count,
                wrote=True,
            )

        existing: dict[str, Any] = {}
        existing_payload: bytes | None = None
        if path.exists():
            existing, existing_payload = _load_json_object(path)

        if key_conflict == "overwrite":
            merged = dict(existing)
            merged.update(incoming)
        elif key_conflict == "error":
            conflicts = sorted(key for key in incoming if key in existing)
            if conflicts:
                names = ", ".join(conflicts)
                raise ServiceError(
                    kind="validation",
                    message=f"JSON key conflict for existing key(s): {names}.",
                    provider="export",
                )
            merged = {**existing, **incoming}
        else:
            merged = dict(existing)
            for key, value in incoming.items():
                if key not in merged:
                    merged[key] = value

        if path.exists() and merged == existing and existing_payload is not None:
            char_count, line_count = _json_stats(existing_payload)
            return TextExportResult(
                path=path,
                char_count=char_count,
                line_count=line_count,
                wrote=False,
            )

        payload = _dump_json_object(merged)
        if path.exists() and existing_payload == payload:
            char_count, line_count = _json_stats(payload)
            return TextExportResult(
                path=path,
                char_count=char_count,
                line_count=line_count,
                wrote=False,
            )

        _atomic_write_with_limit(path, payload)
        char_count, line_count = _json_stats(payload)
        return TextExportResult(
            path=path,
            char_count=char_count,
            line_count=line_count,
            wrote=True,
        )


def writer_for_format(fmt: str) -> ExportWriter:
    """Resolve table writer implementation for a normalized format token."""
    token = fmt.strip().lower()
    if token == "csv":
        return CsvExportWriter()
    if token == "jsonl":
        return JsonlExportWriter()
    if token == "xlsx":
        return XlsxExportWriter()
    raise ServiceError(
        kind="validation",
        message=f"Unsupported export format '{fmt}'.",
        provider="export",
    )


def text_writer_for_format(fmt: str) -> TextExportWriter:
    """Resolve text writer implementation for a normalized format token."""
    token = fmt.strip().lower().lstrip(".")
    if token == "txt":
        return TxtTextExportWriter()
    if token == "docx":
        return DocxTextExportWriter()
    if token == "json":
        return JsonTextExportWriter()
    raise ServiceError(
        kind="validation",
        message=f"Unsupported text export format '{fmt}'.",
        provider="export",
    )
